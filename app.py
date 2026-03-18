import streamlit as st
from fractions import Fraction
from io import BytesIO
from PIL import Image
import json

# 匯入最新的 API 套件
from google import genai

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib as mpl
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
import matplotlib.path as mpath
import matplotlib.patches as mpatches

# LaTeX 風格字體
mpl.rcParams["mathtext.fontset"] = "cm"
mpl.rcParams["font.family"] = "serif"

# ---------------- Unicode 上標轉換器 ----------------
SUPERSCRIPTS = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', 
                '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹', '-': '⁻'}

def to_superscript(n: int):
    return "".join(SUPERSCRIPTS.get(c, c) for c in str(n))

def poly_to_unicode_flow(poly_map, var="x", detached=False):
    if not poly_map: return "0"
    s = ""
    first = True
    for d in sorted(poly_map.keys(), reverse=True):
        a = poly_map.get(d, Fraction(0))
        if a == 0 and not detached: continue
        
        if a.denominator == 1:
            coef_str = str(abs(a.numerator))
        else:
            coef_str = f"{abs(a.numerator)}/{a.denominator}"
            
        if detached:
            body = coef_str
        else:
            if d == 0:
                body = coef_str
            else:
                c_part = "" if abs(a) == 1 else coef_str
                v_part = var if d == 1 else f"{var}{to_superscript(d)}"
                body = f"{c_part}{v_part}"
                
        if first:
            s += f"-{body}" if a < 0 else f"{body}"
            first = False
        else:
            s += f" - {body}" if a < 0 else f" + {body}"
    return s if s else "0"

# ---------------- 多項式工具 ----------------
def parse_coeffs(s: str):
    s = s.replace(",", " ").strip()
    if not s: return [Fraction(0)]
    out = []
    for p in s.split():
        if "/" in p:
            a, b = p.split("/")
            out.append(Fraction(int(a), int(b)))
        else:
            out.append(Fraction(p))
    i = 0
    while i < len(out) - 1 and out[i] == 0: i += 1
    return out[i:] if out else [Fraction(0)]

def strip_leading_zeros(coeffs):
    i = 0
    while i < len(coeffs) - 1 and coeffs[i] == 0: i += 1
    return coeffs[i:]

def poly_to_map(coeffs):
    coeffs = strip_leading_zeros(coeffs)
    n = len(coeffs) - 1
    d = {}
    for i, c in enumerate(coeffs):
        deg = n - i
        d[deg] = d.get(deg, Fraction(0)) + c
    return d

def add_maps(a, b, sgn=1):
    out = dict(a)
    for k, v in b.items():
        out[k] = out.get(k, Fraction(0)) + sgn * v
        if out[k] == 0: del out[k]
    return out

def mul_map_scalar_xk(m, scalar, k):
    out = {}
    for deg, c in m.items(): out[deg + k] = c * scalar
    return out

def leading_term(m):
    if not m: return (None, Fraction(0))
    d = max(m.keys())
    return d, m[d]

def long_division(dividend_coeffs, divisor_coeffs):
    A = poly_to_map(dividend_coeffs)
    B = poly_to_map(divisor_coeffs)
    if not B: raise ValueError("除式不能為 0")
    degB, lcB = leading_term(B)
    if lcB == 0: raise ValueError("除式首項係數為 0")
    Q = {}
    R = dict(A)
    steps = []
    while R and (max(R.keys()) >= degB):
        degR, lcR = leading_term(R)
        k = degR - degB
        t = lcR / lcB
        Q[k] = Q.get(k, Fraction(0)) + t
        P = mul_map_scalar_xk(B, t, k)
        R = add_maps(R, P, sgn=-1)
        steps.append({"product": P, "remainder": dict(R)})
    return Q, steps, R, A, B

# ---------------- mathtext 版面 ----------------
MINUS = "−"

def frac_to_math(fr: Fraction):
    if fr.denominator == 1: return str(fr.numerator)
    return rf"\frac{{{fr.numerator}}}{{{fr.denominator}}}"

def term_to_math(coef_abs: Fraction, deg: int, var: str, detached=False):
    if detached: return frac_to_math(coef_abs) 
    if deg == 0: return frac_to_math(coef_abs)
    base = var if coef_abs == 1 else rf"{frac_to_math(coef_abs)}{var}"
    return base if deg == 1 else rf"{base}^{deg}"

def zero_term_to_math(deg: int, var: str, detached=False):
    if detached: return "0" 
    if deg == 0: return "0"
    return rf"0{var}" if deg == 1 else rf"0{var}^{deg}"

def poly_to_math_flow(poly_map, var, detached=False):
    s = ""
    first = True
    for d in sorted(poly_map.keys(), reverse=True):
        a = poly_map.get(d, Fraction(0))
        if a == 0 and not detached: continue
        body = term_to_math(abs(a), d, var, detached)
        if first:
            s += rf"{MINUS} {body}" if a < 0 else rf"{body}"
            first = False
        else:
            s += rf" {MINUS} {body}" if a < 0 else rf" + {body}"
    return s or "0"

def measure_math(ax, renderer, s, fontsize):
    t = ax.text(0, 0, f"${s}$", fontsize=fontsize, va="baseline")
    bb = t.get_window_extent(renderer=renderer)
    t.remove()
    return bb.width, bb.height

def measure_text(ax, renderer, s, fontsize):
    t = ax.text(0, 0, s, fontsize=fontsize, va="baseline")
    bb = t.get_window_extent(renderer=renderer)
    t.remove()
    return bb.width, bb.height

def draw_math(ax, x, y, s, fontsize, color="black"):
    ax.text(x, y, f"${s}$", fontsize=fontsize, ha="left", va="baseline", color=color)

def draw_text(ax, x, y, s, fontsize, ha="left", va="baseline", color="black"):
    ax.text(x, y, s, fontsize=fontsize, ha=ha, va=va, color=color)

# ======================================================================
# 繪圖核心 Q (題目區)
# ======================================================================
def build_image_math_q(dividend_str, divisor_str, var="x", base_font_size=24, fill_zeros=False, detached_mode=False, question_only=False):
    dividend = parse_coeffs(dividend_str)
    divisor = parse_coeffs(divisor_str)
    Q, steps, R, A, B = long_division(dividend, divisor)

    has_frac = False
    if question_only:
        for M in (A, B):
            if M and any(v.denominator != 1 for v in M.values()):
                has_frac = True; break
    else:
        for M in (A, B, Q):
            if M and any(v.denominator != 1 for v in M.values()):
                has_frac = True; break
        if not has_frac:
            for st in steps:
                if any(v.denominator != 1 for v in st["product"].values()) or \
                   any(v.denominator != 1 for v in st["remainder"].values()):
                    has_frac = True; break

    fs = base_font_size
    margin = 30
    thick = max(2, int(fs * 0.10))
    thin = max(1, int(fs * 0.06))

    if has_frac:
        line_to_div_gap = fs * 1.6
        qu_to_line_gap = fs * 0.9
        line_gap = int(fs * 2.8)
        under_gap = fs * 1.2
    elif not detached_mode:
        line_to_div_gap = fs * 1.1
        qu_to_line_gap = fs * 0.4
        line_gap = int(fs * 1.8)
        under_gap = fs * 0.5
    else:
        line_to_div_gap = fs * 0.8
        qu_to_line_gap = fs * 0.3
        line_gap = int(fs * 1.6)
        under_gap = fs * 0.5

    y_qu = margin + fs
    y_top_line = y_qu + qu_to_line_gap
    y_div = y_top_line + line_to_div_gap

    dpi = 150
    fig = Figure(dpi=dpi)
    canvas = FigureCanvas(fig)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_axis_off() 
    renderer = canvas.get_renderer()

    if question_only:
        all_degs = set(A.keys())
    else:
        all_degs = set()
        for M in (A, B, Q):
            all_degs |= set(M.keys()) if M else set()
        for st in steps:
            all_degs |= set(st["product"].keys())
            all_degs |= set(st["remainder"].keys())
    if not all_degs: all_degs = {0}
    degrees = sorted(all_degs, reverse=True)

    sign_w = max(measure_text(ax, renderer, "+", fs)[0], measure_text(ax, renderer, MINUS, fs)[0])
    sign_gap = 4
    col_w = {d: 0 for d in degrees}
    def consider(poly):
        for d in degrees:
            a = poly.get(d, Fraction(0))
            if a == 0: continue
            w, _ = measure_math(ax, renderer, term_to_math(abs(a), d, var, detached_mode), fs)
            col_w[d] = max(col_w[d], sign_w + sign_gap + w)

    if question_only:
        consider(A)
    else:
        consider(A); consider(B); consider(Q)
        for st in steps:
            consider(st["product"]); consider(st["remainder"])

    if fill_zeros:
        w0, _ = measure_math(ax, renderer, zero_term_to_math(0, var, detached_mode), fs)
        w1, _ = measure_math(ax, renderer, zero_term_to_math(1, var, detached_mode), fs)
        w2, _ = measure_math(ax, renderer, zero_term_to_math(2, var, detached_mode), fs)
        w0max = max(w0, w1, w2)
        for d in degrees: col_w[d] = max(col_w[d], sign_w + sign_gap + w0max)

    div_flow = poly_to_math_flow(B, var, detached_mode)
    div_w, _ = measure_math(ax, renderer, div_flow, fs)

    cols_total = int(sum(col_w[d] for d in degrees))
    cols_gap = 12
    left_gap_px = int(fs * 1.0)
    width = margin + int(div_w) + left_gap_px + cols_gap + cols_total + margin
    
    ax.set_xlim(0, max(10, width))
    ax.set_ylim(max(10, y_div + under_gap * 10), 0)

    x_div_left = margin
    x_divider = x_div_left + int(div_w) + left_gap_px
    x_cols_start = x_divider + cols_gap
    col_left = {}; x_cur = x_cols_start
    for d in degrees:
        col_left[d] = x_cur
        x_cur += int(col_w[d])
    x_cols_end = x_cur

    if not question_only:
        y_products = []; y_mid_lines = []; y_now = y_div
        for _ in steps:
            y_prod = y_now + line_gap
            y_products.append(y_prod)
            y_mid_lines.append(y_prod + under_gap)
            y_now = y_prod + line_gap
        
        y_bottom_line = y_products[-1] + under_gap if steps else y_div + under_gap
        y_rem_final = y_products[-1] + line_gap if steps else y_div + line_gap

    def draw_poly_columns(poly_map, y0, mode='none', lead_deg=None, force_zero_if_all_zero=False):
        first_drawn = True
        for d in degrees:
            a = poly_map.get(d, Fraction(0))
            draw_this = False; body = None
            if mode == 'none' and a != 0:
                draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)
            elif mode == 'all':
                draw_this = True; body = term_to_math(abs(a), d, var, detached_mode) if a != 0 else zero_term_to_math(d, var, detached_mode)
            elif mode == 'window':
                degB, _ = leading_term(B)
                if lead_deg is None or degB is None:
                    if a != 0: draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)
                else:
                    if (lead_deg - degB) <= d <= lead_deg:
                        draw_this = True; body = term_to_math(abs(a), d, var, detached_mode) if a != 0 else zero_term_to_math(d, var, detached_mode)
                    elif a != 0:
                        draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)
            elif mode == 'pad_to_zero': 
                if lead_deg is not None and 0 <= d <= lead_deg:
                    draw_this = True; body = term_to_math(abs(a), d, var, detached_mode) if a != 0 else zero_term_to_math(d, var, detached_mode)
                elif a != 0:
                    draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)

            if not draw_this: continue
            cx = col_left[d]
            
            if first_drawn:
                if a < 0: draw_text(ax, cx, y0, MINUS, fs)
                draw_math(ax, cx + sign_w + sign_gap, y0, body, fs)
                first_drawn = False
            else:
                draw_text(ax, cx, y0, MINUS if a < 0 else "+", fs)
                draw_math(ax, cx + sign_w + sign_gap, y0, body, fs)
                
        if first_drawn and force_zero_if_all_zero:
            d0 = 0 if 0 in degrees else degrees[-1]
            draw_math(ax, col_left[d0] + sign_w + sign_gap, y0, "0", fs)

    if not question_only:
        draw_poly_columns(Q, y_qu, mode='none')
    
    if has_frac:
        y_bracket_bottom = y_div + fs * 0.5
    else:
        y_bracket_bottom = y_div + fs * 0.2
        
    curve_w = fs * 0.35
    H = y_bracket_bottom - y_top_line
    
    verts = [
        (x_divider - curve_w, y_top_line),                       
        (x_divider + curve_w * 0.2, y_top_line + H * 0.2),       
        (x_divider + curve_w * 0.2, y_bracket_bottom - H * 0.2), 
        (x_divider - curve_w, y_bracket_bottom)                  
    ]
    codes = [mpath.Path.MOVETO, mpath.Path.CURVE4, mpath.Path.CURVE4, mpath.Path.CURVE4]
    path = mpath.Path(verts, codes)
    patch = mpatches.PathPatch(path, facecolor='none', edgecolor='black', lw=thick)
    ax.add_patch(patch)
    
    ax.plot([x_divider - curve_w, x_cols_end], [y_top_line, y_top_line], color="black", linewidth=thick)

    draw_math(ax, x_div_left, y_div, poly_to_math_flow(B, var, detached_mode), fs)
    draw_poly_columns(A, y_div, mode=('all' if fill_zeros else 'none'))

    if not question_only:
        for i, st in enumerate(steps):
            y_prod = y_products[i]
            draw_poly_columns(st["product"], y_prod, mode='none')
            ax.plot([x_cols_start, x_cols_end], [y_mid_lines[i], y_mid_lines[i]], color="black", linewidth=thin)
            is_last = (i == len(steps) - 1)
            y_rem = y_prod + line_gap
            if fill_zeros and not is_last:
                lead_deg, _ = leading_term(st["remainder"])
                draw_poly_columns(st["remainder"], y_rem, mode='window', lead_deg=lead_deg)
            else:
                draw_poly_columns(st["remainder"], y_rem, mode='none')

        ax.plot([x_cols_start, x_cols_end], [y_bottom_line, y_bottom_line], color="black", linewidth=thick)
        final_R = R if R else (steps[-1]["remainder"] if steps else A)
        
        if fill_zeros:
            final_lead = max(final_R.keys()) if final_R else 0
            draw_poly_columns(final_R, y_rem_final, mode='pad_to_zero', lead_deg=final_lead, force_zero_if_all_zero=True)
        else:
            draw_poly_columns(final_R, y_rem_final, mode='none', force_zero_if_all_zero=True)

    canvas.draw()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, facecolor="white", bbox_inches="tight", pad_inches=0.05)
    buf.seek(0)
    plt.close(fig)
    return Image.open(buf)

# ======================================================================
# 繪圖核心 A (解答區)
# ======================================================================
def build_image_math_a(dividend_str, divisor_str, var="x", base_font_size=24, fill_zeros=False, detached_mode=False, question_only=False):
    dividend = parse_coeffs(dividend_str)
    divisor = parse_coeffs(divisor_str)
    Q, steps, R, A, B = long_division(dividend, divisor)
    degB, lcB = leading_term(B)

    has_frac = False
    if question_only:
        for M in (A, B):
            if M and any(v.denominator != 1 for v in M.values()):
                has_frac = True; break
    else:
        for M in (A, B, Q):
            if M and any(v.denominator != 1 for v in M.values()):
                has_frac = True; break
        if not has_frac:
            for st in steps:
                if any(v.denominator != 1 for v in st["product"].values()) or \
                   any(v.denominator != 1 for v in st["remainder"].values()):
                    has_frac = True; break

    fs = base_font_size
    margin = 30
    y_qu = margin + fs
    
    if has_frac:
        line_gap = int(fs * 3.6)       
        under_gap = fs * 1.5           
        y_top_line = y_qu + fs * 1.8   
        y_div = y_top_line + fs * 2.4  
    else:
        line_gap = int(fs * 2.2)
        under_gap = fs * 0.7
        y_top_line = y_qu + fs * 0.9   
        y_div = y_top_line + fs * 1.6  

    if question_only:
        y_top_line = margin + fs * 0.2
        if has_frac:
            y_div = y_top_line + fs * 1.5
        else:
            y_div = y_top_line + fs * 1.1

    dpi = 150
    fig = Figure(dpi=dpi)
    canvas = FigureCanvas(fig)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_axis_off()
    renderer = canvas.get_renderer()

    thin = max(1, int(fs * 0.06))
    thick = max(2, int(fs * 0.10))

    if question_only:
        all_degs = set(A.keys()) | set(B.keys())
    else:
        all_degs = set()
        for M in (A, B, Q):
            all_degs |= set(M.keys()) if M else set()
        for st in steps:
            all_degs |= set(st["product"].keys())
            all_degs |= set(st["remainder"].keys())
            
    if not all_degs: all_degs = {0}
    degrees = sorted(all_degs, reverse=True)

    sign_w = max(measure_text(ax, renderer, "+", fs)[0], measure_text(ax, renderer, MINUS, fs)[0])
    sign_gap = 4
    col_w = {d: 0 for d in degrees}
    
    def consider(poly):
        for d in degrees:
            a = poly.get(d, Fraction(0))
            if a == 0: continue
            w, _ = measure_math(ax, renderer, term_to_math(abs(a), d, var, detached_mode), fs)
            col_w[d] = max(col_w[d], sign_w + sign_gap + w)
            
    if question_only:
        consider(A); consider(B)
    else:
        consider(A); consider(B); consider(Q)
        for st in steps:
            consider(st["product"]); consider(st["remainder"])

    if fill_zeros:
        w0, _ = measure_math(ax, renderer, zero_term_to_math(0, var, detached_mode), fs)
        w1, _ = measure_math(ax, renderer, zero_term_to_math(1, var, detached_mode), fs)
        w2, _ = measure_math(ax, renderer, zero_term_to_math(2, var, detached_mode), fs)
        w0max = max(w0, w1, w2)
        for d in degrees: col_w[d] = max(col_w[d], sign_w + sign_gap + w0max)

    div_flow = poly_to_math_flow(B, var, detached_mode)
    div_w, _ = measure_math(ax, renderer, div_flow, fs)

    cols_total = int(sum(col_w[d] for d in degrees))
    cols_gap = 12
    left_gap_px = int(fs * 1.2)
    width = margin + int(div_w) + left_gap_px + cols_gap + cols_total + margin
    s_len = len(steps)
    
    if question_only:
        total_rows = 2
        height = y_div + under_gap * 1.5 + margin  
    else:
        total_rows = (2 * s_len + 2) if s_len > 0 else 3
        height = margin + total_rows * line_gap + margin
        
    ax.set_xlim(0, max(10, width))
    ax.set_ylim(max(10, height), 0)

    x_div_left = margin
    x_divider = x_div_left + int(div_w) + left_gap_px
    x_cols_start = x_divider + cols_gap
    col_left = {}; x_cur = x_cols_start
    for d in degrees:
        col_left[d] = x_cur
        x_cur += int(col_w[d])
    x_cols_end = x_cur

    if not question_only:
        y_products = []; y_mid_lines = []; y_now = y_div
        for _ in steps:
            y_prod = y_now + line_gap
            y_products.append(y_prod)
            y_mid_lines.append(y_prod + under_gap)
            y_now = y_prod + line_gap
        
        y_bottom_line = y_products[-1] + under_gap if steps else y_div + under_gap
        y_rem_final = y_products[-1] + line_gap if steps else y_div + line_gap

    def draw_poly_columns(poly_map, y0, mode='none', lead_deg=None, force_zero_if_all_zero=False):
        first_drawn = True
        for d in degrees:
            a = poly_map.get(d, Fraction(0))
            draw_this = False; body = None
            if mode == 'none' and a != 0:
                draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)
            elif mode == 'all':
                draw_this = True; body = term_to_math(abs(a), d, var, detached_mode) if a != 0 else zero_term_to_math(d, var, detached_mode)
            elif mode == 'window':
                if lead_deg is None or degB is None:
                    if a != 0: draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)
                else:
                    if (lead_deg - degB) <= d <= lead_deg:
                        draw_this = True; body = term_to_math(abs(a), d, var, detached_mode) if a != 0 else zero_term_to_math(d, var, detached_mode)
                    elif a != 0:
                        draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)
            elif mode == 'pad_to_zero': 
                if lead_deg is not None and 0 <= d <= lead_deg:
                    draw_this = True; body = term_to_math(abs(a), d, var, detached_mode) if a != 0 else zero_term_to_math(d, var, detached_mode)
                elif a != 0:
                    draw_this = True; body = term_to_math(abs(a), d, var, detached_mode)

            if not draw_this: continue
            cx = col_left[d]
            
            if first_drawn:
                if a < 0: draw_text(ax, cx, y0, MINUS, fs)
                draw_math(ax, cx + sign_w + sign_gap, y0, body, fs)
                first_drawn = False
            else:
                draw_text(ax, cx, y0, MINUS if a < 0 else "+", fs)
                draw_math(ax, cx + sign_w + sign_gap, y0, body, fs)
                
        if first_drawn and force_zero_if_all_zero:
            d0 = 0 if 0 in degrees else degrees[-1]
            draw_math(ax, col_left[d0] + sign_w + sign_gap, y0, "0", fs)

    if not question_only:
        draw_poly_columns(Q, y_qu, mode='none')
    
    curve_w = fs * 0.4
    y_bracket_bottom = y_div + under_gap * 0.8
    H = y_bracket_bottom - y_top_line
    
    verts = [
        (x_divider - curve_w, y_top_line),                       
        (x_divider + curve_w * 0.1, y_top_line + H * 0.2),       
        (x_divider + curve_w * 0.1, y_bracket_bottom - H * 0.2), 
        (x_divider - curve_w, y_bracket_bottom)                  
    ]
    codes = [mpath.Path.MOVETO, mpath.Path.CURVE4, mpath.Path.CURVE4, mpath.Path.CURVE4]
    path = mpath.Path(verts, codes)
    patch = mpatches.PathPatch(path, facecolor='none', edgecolor='black', lw=thick)
    ax.add_patch(patch)
    ax.plot([x_divider - curve_w, x_cols_end], [y_top_line, y_top_line], color="black", linewidth=thick)

    draw_math(ax, x_div_left, y_div, poly_to_math_flow(B, var, detached_mode), fs)
    draw_poly_columns(A, y_div, mode=('all' if fill_zeros else 'none'))

    if not question_only:
        for i, st in enumerate(steps):
            y_prod = y_products[i]
            draw_poly_columns(st["product"], y_prod, mode='none')
            ax.plot([x_cols_start, x_cols_end], [y_mid_lines[i], y_mid_lines[i]], color="black", linewidth=thin)
            is_last = (i == len(steps) - 1)
            y_rem = y_prod + line_gap
            if fill_zeros and not is_last:
                lead_deg, _ = leading_term(st["remainder"])
                draw_poly_columns(st["remainder"], y_rem, mode='window', lead_deg=lead_deg)
            else:
                draw_poly_columns(st["remainder"], y_rem, mode='none')

        ax.plot([x_cols_start, x_cols_end], [y_bottom_line, y_bottom_line], color="black", linewidth=thick)
        final_R = R if R else (steps[-1]["remainder"] if steps else A)
        
        if fill_zeros:
            final_lead = max(final_R.keys()) if final_R else 0
            draw_poly_columns(final_R, y_rem_final, mode='pad_to_zero', lead_deg=final_lead, force_zero_if_all_zero=True)
        else:
            draw_poly_columns(final_R, y_rem_final, mode='none', force_zero_if_all_zero=True)

    canvas.draw()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, facecolor="white", bbox_inches="tight", pad_inches=0.15)
    buf.seek(0)
    plt.close(fig)
    return Image.open(buf)

# ======================================================================
# 路由邏輯
# ======================================================================
def build_image_math(dividend_str, divisor_str, var="x", base_font_size=24, fill_zeros=False, detached_mode=False, question_only=False):
    if question_only:
        return build_image_math_q(dividend_str, divisor_str, var, base_font_size, fill_zeros, detached_mode, question_only)
    else:
        return build_image_math_a(dividend_str, divisor_str, var, base_font_size, fill_zeros, detached_mode, question_only)

# ======================================================================
# Streamlit 網頁應用程式主體
# ======================================================================
st.set_page_config(page_title="阿凱老師的長除法產生器", page_icon="👨‍🏫", layout="wide")

# 初始化 session_state
if "dividend" not in st.session_state:
    st.session_state.dividend = "1, 2, 1, -8"
if "divisor" not in st.session_state:
    st.session_state.divisor = "1, -2"

st.title("👨‍🏫 阿凱老師的長除法產生器")
st.markdown("專為國中數學老師打造！一鍵產生精美排版的「多項式長除法」題目與解答圖片，並支援 AI 出題與匯出 Word 學習單。")

# 側邊欄設定
with st.sidebar:
    st.header("⚙️ 參數設定")
    var_val = st.text_input("變數符號", value="x")
    
    # 綁定 session_state
    dividend_val = st.text_input("被除式係數 (高次到低次，逗號分隔)", key="dividend")
    divisor_val = st.text_input("除式係數 (高次到低次，逗號分隔)", key="divisor")

    st.markdown("---")
    st.subheader("🎨 排版選項")
    fill_zeros = st.checkbox("缺項補零", value=True)
    detached_mode = st.checkbox("分離係數法 (不顯示變數)", value=False)
    
    st.markdown("---")
    st.subheader("✨ AI 自動出題 (Gemini)")
    api_key = st.text_input("輸入您的 Gemini API Key", type="password")
    
    # AI 難度自訂
    col1, col2 = st.columns(2)
    with col1:
        deg_div = st.selectbox("被除式最高次", [2, 3, 4, 5], index=1)
    with col2:
        deg_dsr = st.selectbox("除式最高次", [1, 2, 3], index=1)

    def generate_one_question():
        if not api_key:
            st.error("⚠️ 請先輸入 API Key！")
            return
        try:
            client = genai.Client(api_key=api_key)
            prompt = f"""
            請產生 1 個適合國中數學的「多項式長除法」計算題。
            請嚴格以 JSON 格式輸出，不要有任何 Markdown 或說明文字：
            {{"dividend": "被除式的係數", "divisor": "除式的係數"}}
            要求：
            1. 被除式最高次「必須剛好是」 {deg_div} 次。
            2. 除式最高次「必須剛好是」 {deg_dsr} 次。
            3. 係數請用逗號分隔，例如 "1, -2, 0, 5"。
            4. 適合手算，數字不要太大。
            """
            response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
            raw_text = response.text.replace('```json', '').replace('```', '').strip()
            data = json.loads(raw_text)
            
            # 更新 session_state，Streamlit 會自動重新整理畫面
            st.session_state.dividend = data.get("dividend", "")
            st.session_state.divisor = data.get("divisor", "")
            st.success("✅ 題目產生成功！")
        except Exception as e:
            st.error(f"❌ 錯誤: {e}")

    st.button("🎲 AI 產生 1 題 (更新畫面)", on_click=generate_one_question, use_container_width=True)
    
    st.markdown("---")
    st.subheader("📄 匯出學習單")
    
    # Word 生成按鈕區塊
    if st.button("📦 產生 5 題並準備下載", type="primary", use_container_width=True):
        if not api_key:
            st.error("⚠️ 請先輸入 API Key！")
        else:
            with st.spinner("⏳ 正在請 AI 出 5 題，並排版成 Word，請稍候..."):
                try:
                    import docx
                    from docx.shared import Inches, Pt
                    
                    client = genai.Client(api_key=api_key)
                    prompt = f"""
                    請產生 5 個不同的國中數學「多項式長除法」計算題。
                    請嚴格以 JSON 陣列 (Array) 格式輸出，不要有任何 Markdown 或說明文字。格式如下：
                    [
                        {{"dividend": "被除式的係數", "divisor": "除式的係數"}},
                        ... (請確保剛好 5 個物件)
                    ]
                    要求：
                    1. 題目難度適合國中生手算。
                    2. 被除式最高次「必須剛好是」 {deg_div} 次。
                    3. 除式最高次「必須剛好是」 {deg_dsr} 次。
                    4. 係數請用逗號分隔，例如 "2, 3, -5"。
                    """
                    response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
                    raw_text = response.text.replace('```json', '').replace('```', '').strip()
                    data = json.loads(raw_text)
                    
                    doc = docx.Document()
                    doc.add_heading('多項式長除法 練習卷', 0)
                    
                    # 產生題目區
                    doc.add_heading('一、 計算題', level=1)
                    for i, q in enumerate(data):
                        div_coeffs = parse_coeffs(q["dividend"])
                        divisor_coeffs = parse_coeffs(q["divisor"])
                        A_poly = poly_to_map(div_coeffs)
                        B_poly = poly_to_map(divisor_coeffs)
                        A_uni = poly_to_unicode_flow(A_poly, var_val, detached=False)
                        B_uni = poly_to_unicode_flow(B_poly, var_val, detached=False)
                        
                        p = doc.add_paragraph(f"第 {i+1} 題：計算 ( {A_uni} ) ÷ ( {B_uni} ) 的商式與餘式。")
                        p.runs[0].font.size = Pt(12)
                        
                        # 匯出 Word 圖片字級為 30
                        img_q = build_image_math(
                            q["dividend"], q["divisor"], var=var_val, base_font_size=30,
                            fill_zeros=fill_zeros, detached_mode=detached_mode, question_only=True
                        )
                        img_byte_arr = BytesIO()
                        img_q.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        doc.add_picture(img_byte_arr, width=Inches(3.5))
                        doc.add_paragraph('\n\n') 
                        
                    doc.add_page_break()
                    
                    # 產生解答區
                    doc.add_heading('解答區', level=1)
                    for i, q in enumerate(data):
                        doc.add_paragraph(f"第 {i+1} 題 解答：")
                        
                        img_a = build_image_math(
                            q["dividend"], q["divisor"], var=var_val, base_font_size=30,
                            fill_zeros=fill_zeros, detached_mode=detached_mode, question_only=False
                        )
                        img_byte_arr = BytesIO()
                        img_a.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        doc.add_picture(img_byte_arr, width=Inches(4.5))
                        doc.add_paragraph('\n')
                        
                    # 存入 BytesIO 以供下載
                    bio = BytesIO()
                    doc.save(bio)
                    st.session_state.docx_data = bio.getvalue()
                    st.success("✅ Word 學習單已準備就緒！請點擊下方按鈕下載。")

                except Exception as e:
                    st.error(f"❌ 錯誤: {e}")

    # 若有產生過 Word，顯示下載按鈕
    if "docx_data" in st.session_state:
        st.download_button(
            label="📥 點我下載 Word 學習單",
            data=st.session_state.docx_data,
            file_name="多項式長除法_學習單.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# 主畫面渲染區
try:
    A_poly = poly_to_map(parse_coeffs(st.session_state.dividend))
    B_poly = poly_to_map(parse_coeffs(st.session_state.divisor))
    
    A_uni = poly_to_unicode_flow(A_poly, var_val, detached=False)
    B_uni = poly_to_unicode_flow(B_poly, var_val, detached=False)
    
    st.info(f"### 📝 題目：試計算 **( {A_uni} ) ÷ ( {B_uni} )** 的商式與餘式。")
    
    col_q, col_a = st.columns(2)
    
    with col_q:
        st.markdown("#### 📝 題目區 (緊湊版)")
        img_q = build_image_math(
            st.session_state.dividend, st.session_state.divisor, var=var_val, 
            fill_zeros=fill_zeros, detached_mode=detached_mode, question_only=True
        )
        st.image(img_q, use_container_width=False)
        
    with col_a:
        st.markdown("#### 💡 解答區")
        img_a = build_image_math(
            st.session_state.dividend, st.session_state.divisor, var=var_val, 
            fill_zeros=fill_zeros, detached_mode=detached_mode, question_only=False
        )
        st.image(img_a, use_container_width=False)

except Exception as e:
    st.error(f"輸入格式有誤或無法計算: {e}")